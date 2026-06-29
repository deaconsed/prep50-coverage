"""Convert a WAEC objective-paper DOCX to a Prep50-ready CSV.

Each paragraph in the source DOCX is expected to be one complete question:

    <stem> A. <opt1> B. <opt2> C. <opt3> D. <opt4>[.] [YYYY/N]

Output CSV columns: question, question_year, option_1, option_2, option_3,
option_4, short_answer. short_answer is left blank (the original DOCX doesn't
carry it); question_year defaults to whatever --year says (2026 unless overridden).

Usage:
    python scripts/docx_to_csv.py --input "WAEC CRS OBJ (2).docx" --out crs_2026.csv
    python scripts/docx_to_csv.py --input paper.docx --out paper.csv --year 2026
"""
from __future__ import annotations

import argparse
import csv
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

# --- Math / image-aware text extraction -----------------------------------
# python-docx's Paragraph.text and _Cell.text only read plain text runs. They
# silently DROP equations (OMML <m:oMath> objects) and images, and they FLATTEN
# subscript/superscript runs to baseline (so H<sub>2</sub>O -> "H2O", x<sup>2</sup>
# -> "x2", losing the structure that distinguishes a question). That guts the
# signal the embedder relies on for equation-heavy papers. The helpers below
# replace .text everywhere so that:
#   - OMML equations are linearised to readable math (frac -> (a)/(b),
#     superscript -> ^, subscript -> _, radical -> sqrt(), n-ary -> operator).
#   - subscript/superscript runs are marked with _ / ^ instead of flattened.
#   - images emit an [IMAGE] placeholder so a figure-only option/stem is kept
#     and visible (route these through OCR later) instead of vanishing.
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
IMAGE_PLACEHOLDER = " [IMAGE] "


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _m(tag: str) -> str:
    return f"{{{M_NS}}}{tag}"


def _supwrap(s: str) -> str:
    """Brace multi-character exponents/indices so 'x^-2' reads as 'x^(-2)'."""
    s = s.strip()
    return s if len(s) <= 1 else f"({s})"


def _omml_children(el) -> str:
    if el is None:
        return ""
    return "".join(_omml_text(c) for c in el)


def _omml_group(el, name: str) -> str:
    """Linearised text of a named OMML child container (m:e, m:num, m:sup, …)."""
    child = el.find(_m(name))
    return _omml_children(child) if child is not None else ""


def _omml_text(el) -> str:
    """Recursively linearise an OMML element to a compact ASCII math string."""
    tag = etree.QName(el).localname
    if tag == "t":                                   # math run text
        return el.text or ""
    if tag == "r":                                   # math run
        return "".join(_omml_text(c) for c in el if etree.QName(c).localname == "t")
    if tag == "f":                                   # fraction
        return f"({_omml_group(el, 'num')})/({_omml_group(el, 'den')})"
    if tag == "sSup":                                # superscript  base^exp
        return f"{_omml_group(el, 'e')}^{_supwrap(_omml_group(el, 'sup'))}"
    if tag == "sSub":                                # subscript    base_idx
        return f"{_omml_group(el, 'e')}_{_supwrap(_omml_group(el, 'sub'))}"
    if tag == "sSubSup":                             # base_idx^exp
        return (f"{_omml_group(el, 'e')}"
                f"_{_supwrap(_omml_group(el, 'sub'))}"
                f"^{_supwrap(_omml_group(el, 'sup'))}")
    if tag == "rad":                                 # radical
        deg = _omml_group(el, "deg")
        e = _omml_group(el, "e")
        return f"sqrt({e})" if not deg else f"({e})^(1/{deg})"
    if tag == "d":                                   # delimiter (…)
        return f"({_omml_group(el, 'e')})"
    if tag == "func":                                # function apply  f(x)
        return f"{_omml_group(el, 'fName')}({_omml_group(el, 'e')})"
    if tag == "nary":                                # sum / integral / product
        ch = "∫"
        pr = el.find(_m("naryPr"))
        if pr is not None:
            chr_el = pr.find(_m("chr"))
            if chr_el is not None:
                ch = chr_el.get(_m("val")) or ch
        out = ch
        sub = _omml_group(el, "sub")
        sup = _omml_group(el, "sup")
        if sub:
            out += f"_{_supwrap(sub)}"
        if sup:
            out += f"^{_supwrap(sup)}"
        return f"{out} {_omml_group(el, 'e')}"
    # Containers (m:e, m:num, m:oMath, …) and anything unrecognised: recurse so
    # raw symbols/numbers are still recovered even if the structure isn't modelled.
    return _omml_children(el)


def _render_run(r) -> str:
    """Plain text of a w:r, marking sub/superscript and noting embedded images."""
    txt = "".join(t.text or "" for t in r.findall(_w("t")))
    if txt:
        rpr = r.find(_w("rPr"))
        if rpr is not None:
            va = rpr.find(_w("vertAlign"))
            if va is not None:
                val = va.get(_w("val"))
                if val == "subscript":
                    txt = "_" + _supwrap(txt)
                elif val == "superscript":
                    txt = "^" + _supwrap(txt)
    if any(r.find(_w(t)) is not None for t in ("drawing", "pict", "object")):
        txt += IMAGE_PLACEHOLDER
    return txt


def _walk(el, out: list) -> None:
    qname = etree.QName(el)
    ns, tag = qname.namespace, qname.localname
    if ns == M_NS and tag in ("oMath", "oMathPara"):
        out.append(f" {_omml_text(el)} ")
        return                                       # don't descend into math
    if ns == W_NS:
        if tag == "r":
            out.append(_render_run(el))
            return                                   # run handled whole
        if tag in ("drawing", "pict", "object"):
            out.append(IMAGE_PLACEHOLDER)
            return
        if tag in ("tab", "br", "cr"):
            out.append(" ")
            return
    for child in el:                                 # hyperlink, smartTag, sdt, …
        _walk(child, out)


def para_text(paragraph: Paragraph) -> str:
    """Replacement for Paragraph.text that preserves math and flags images."""
    out: list[str] = []
    _walk(paragraph._p, out)
    return "".join(out)


def cell_text(cell) -> str:
    return " ".join(para_text(p) for p in cell.paragraphs).strip()

# One regex matches the whole question. `(?P<stem>.+?)` is non-greedy and
# bounded by the next required `\s+A\.`, so each option capture stops at the
# next option marker. The space after each marker dot is `\s*` (not `\s+`) so
# that glued markers like "C.one" still match, but each option must then start
# with a non-space char (`\S`) so blank options (equations dropped during text
# extraction) are still rejected. The final `\[YYYY/N\]` tail is required, so
# noise paragraphs (page headers, instructions) get skipped.
QUESTION_RE = re.compile(
    r"""^
        (?P<stem>.+?)
        \s+A\.\s*(?P<a>\S.*?)
        \s+B\.\s*(?P<b>\S.*?)
        \s+C\.\s*(?P<c>\S.*?)
        \s+D\.\s*(?P<d>\S.*?)
        \s*\[\d{4}/\d+\]\s*$
    """,
    re.VERBOSE | re.DOTALL,
)

# Fallback signals for questions whose options didn't survive text extraction
# (formula choices stored as equation objects, or diagram choices stored as
# images). A trailing `[YYYY/N]` tag is the reliable question signature, so any
# paragraph carrying one is kept even when QUESTION_RE can't find four options.
YEAR_TAIL_RE = re.compile(r"\s*\[\d{4}/\d+\]\s*$")
# Strips a trailing run of empty option markers ("A. B. C. D.") off the stem so
# the leftover artifact doesn't pollute the question text.
EMPTY_OPTIONS_RE = re.compile(r"\s+A\.\s*B\.\s*C\.\s*D\.\s*\.?\s*$")

# An [IMAGE] placeholder (from para_text) for a diagram/equation lands in document
# order — usually AFTER the `[YYYY/N]` tag — which would otherwise break the
# end-anchored year-tag match. Peel trailing image markers off before matching and
# re-flag the question as carrying a figure.
TRAILING_IMAGE_RE = re.compile(r"(?:\s*\[IMAGE\])+\s*$")

# Same shape as QUESTION_RE but with no trailing year tag: a last-resort recovery
# of a cleanly 4-optioned question whose `[YYYY/N]` tag is missing in the source.
QUESTION_NO_YEAR_RE = re.compile(
    r"""^
        (?P<stem>.+?)
        \s+A\.\s*(?P<a>\S.*?)
        \s+B\.\s*(?P<b>\S.*?)
        \s+C\.\s*(?P<c>\S.*?)
        \s+D\.\s*(?P<d>\S.*?)
        \s*$
    """,
    re.VERBOSE | re.DOTALL,
)

# WAEC groups several questions under a block of shared data introduced by a
# "Use the following information to answer question(s) …" header. The data
# itself is a mix of continuation paragraphs and tables, and is reprinted before
# each question in the group. We detect the header, accumulate the data, and
# prepend it to the next question's stem so the question stays self-contained.
CONTEXT_HEADER_RE = re.compile(r"^Use the following information to answer question", re.I)

# Cyrillic and Greek homoglyphs sometimes leak in from OCR/copy-paste of WAEC
# papers. They render identically to ASCII but break the option-marker regex
# (e.g. a Greek capital Beta "Β" standing in for the "B." option marker).
_HOMOGLYPHS = str.maketrans({
    # Cyrillic
    "А": "A", "В": "B", "С": "C", "Е": "E", "Н": "H", "К": "K",
    "М": "M", "О": "O", "Р": "P", "Т": "T", "Х": "X",
    "а": "a", "е": "e", "о": "o", "р": "p", "с": "c",
    "у": "y", "х": "x",
    # Greek capitals that are visually identical to Latin letters. Restricted
    # to true look-alikes so genuine Greek physics symbols (Ω, Δ, Σ, …) are
    # left untouched.
    "Α": "A", "Β": "B", "Ε": "E", "Ζ": "Z", "Η": "H", "Ι": "I",
    "Κ": "K", "Μ": "M", "Ν": "N", "Ο": "O", "Ρ": "P", "Τ": "T",
    "Υ": "Y", "Χ": "X",
})


def clean_text(s: str) -> str:
    """Collapse whitespace, strip stray leading punctuation, normalize quotes."""
    s = re.sub(r"\s+", " ", s).strip()
    s = re.sub(r"^[.\s]+", "", s)               # leading "." or whitespace
    s = re.sub(r"\s+([.,;:?!])", r"\1", s)      # " ?" → "?"
    return s


def clean_option(s: str) -> str:
    """Same as clean_text plus drop a trailing terminal punctuation if any."""
    s = clean_text(s)
    s = re.sub(r"[.,;:]+$", "", s).strip()
    return s


def iter_block_items(doc):
    """Yield Paragraphs and Tables in true document order.

    doc.paragraphs and doc.tables are separate flat lists, so on their own they
    lose the interleaving between a question's preamble paragraphs and the data
    table that belongs with it. Walking the body's child elements preserves it.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def render_table(table: Table) -> str:
    """Flatten a data table to one readable line: rows joined by '; '.

    Merged cells make python-docx repeat a cell's text across the span, so
    duplicates within a row are collapsed.
    """
    rows = []
    for row in table.rows:
        cells = [cell_text(c) for c in row.cells]
        cells = [c for c in dict.fromkeys(cells) if c]  # de-dupe merges, drop blanks
        if cells:
            rows.append(" ".join(cells))
    return "; ".join(rows)


def parse_paragraph(text: str) -> dict | None:
    # Peel any trailing [IMAGE] markers so the end-anchored year tag still matches;
    # remember it so the parsed question keeps a "[IMAGE]" flag on its stem.
    core, had_image = text, False
    if TRAILING_IMAGE_RE.search(text):
        core = TRAILING_IMAGE_RE.sub("", text).rstrip()
        had_image = True

    def flag(stem: str) -> str:
        stem = clean_text(stem)
        return f"{stem} [IMAGE]" if had_image else stem

    def full(g: dict) -> dict:
        return {
            "question": flag(g["stem"]),
            "option_1": clean_option(g["a"]),
            "option_2": clean_option(g["b"]),
            "option_3": clean_option(g["c"]),
            "option_4": clean_option(g["d"]),
        }

    m = QUESTION_RE.match(core)
    if m:
        return full(m.groupdict())

    # Fallback: keep the question even without options, as long as it carries a
    # year tag. The stem is everything before the tag, minus any leftover empty
    # option markers; the four option columns are written blank.
    if YEAR_TAIL_RE.search(core):
        stem = EMPTY_OPTIONS_RE.sub("", YEAR_TAIL_RE.sub("", core))
        return {"question": flag(stem), "option_1": "", "option_2": "",
                "option_3": "", "option_4": ""}

    # Last resort: a cleanly 4-optioned question whose year tag is missing.
    m = QUESTION_NO_YEAR_RE.match(core)
    if m:
        return full(m.groupdict())
    return None


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--input", required=True, type=Path, help="DOCX file to read.")
    ap.add_argument("--out", required=True, type=Path, help="Output CSV path.")
    ap.add_argument("--year", type=int, default=2026,
                    help="question_year for every row (default 2026).")
    args = ap.parse_args()

    if not args.input.exists():
        sys.exit(f"Input not found: {args.input}")

    doc = Document(str(args.input))
    parsed: list[dict] = []
    skipped: list[tuple[int, str]] = []

    context = ""        # shared stimulus accumulated for the current group
    in_context = False  # armed by a "Use the following information" header
    with_context = 0

    for i, item in enumerate(iter_block_items(doc)):
        if isinstance(item, Table):
            # Only fold in tables that belong to an active stimulus block.
            if in_context:
                table_text = render_table(item)
                if table_text:
                    context = f"{context} {table_text}".strip()
            continue

        text = para_text(item).strip().translate(_HOMOGLYPHS)
        if not text:
            continue

        if CONTEXT_HEADER_RE.match(text):
            # New stimulus block (reprinted before each question in the group):
            # start fresh and drop the navigational header sentence itself.
            context = ""
            in_context = True
            continue

        row = parse_paragraph(text)
        if row is None:
            if in_context:
                # A data/continuation line inside the stimulus block.
                context = f"{context} {text}".strip()
            else:
                skipped.append((i, text[:90]))
            continue

        if context:
            row["question"] = clean_text(f"{context} {row['question']}")
            with_context += 1
        row["question_year"] = args.year
        row["short_answer"] = ""
        parsed.append(row)
        context = ""
        in_context = False

    no_options = sum(1 for r in parsed if not r["option_1"])
    print(f"Parsed {len(parsed)} questions.")
    if with_context:
        print(f"  ({with_context} had shared data/preamble folded into the stem.)")
    if no_options:
        print(f"  ({no_options} kept with empty options — choices not in text.)")
    if skipped:
        print(f"Skipped {len(skipped)} paragraph(s) that didn't match:")
        for i, sample in skipped[:5]:
            print(f"  [{i}] {sample!r}…")
        if len(skipped) > 5:
            print(f"  …and {len(skipped) - 5} more")

    cols = [
        "question", "question_year",
        "option_1", "option_2", "option_3", "option_4",
        "short_answer",
    ]
    args.out.parent.mkdir(parents=True, exist_ok=True)
    # utf-8-sig writes a BOM so Excel opens the CSV without mangling smart quotes.
    with open(args.out, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=cols)
        w.writeheader()
        for row in parsed:
            w.writerow(row)
    print(f"Wrote: {args.out}")


if __name__ == "__main__":
    main()
